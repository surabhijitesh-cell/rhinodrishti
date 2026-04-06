from fastapi import FastAPI, APIRouter, Query, HTTPException, BackgroundTasks, UploadFile, File, WebSocket, WebSocketDisconnect
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict
import uuid
import asyncio
from datetime import datetime, timezone, timedelta
import io

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

app = FastAPI(title="Rhino Drishti API")
api_router = APIRouter(prefix="/api")

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

intelligence_col = db.intelligence_items
briefs_col = db.daily_briefs
sources_col = db.rss_sources
uploads_col = db.uploaded_documents
tweets_col = db.twitter_feeds
national_news_col = db.national_news
international_news_col = db.international_news
patterns_col = db.intelligence_patterns

# In-memory scan status tracker
scan_status = {
    "is_scanning": False,
    "progress": 0,
    "total_sources": 0,
    "current_source": "",
    "sources_scanned": 0,
    "articles_found": 0,
    "relevant_found": 0,
    "filtered_out": 0,
    "translated": 0,
    "last_scan_at": None,
    "last_scan_result": None,
    "scan_log": [],
}

# Twitter/X accounts to monitor for defense updates
TWITTER_ACCOUNTS_TO_MONITOR = [
    {"handle": "@adgpi", "name": "ADG PI - Indian Army", "category": "defense", "url": "https://twitter.com/adgpi"},
    {"handle": "@IAF_MCC", "name": "Indian Air Force", "category": "defense", "url": "https://twitter.com/IAF_MCC"},
    {"handle": "@indiannavy", "name": "Indian Navy", "category": "defense", "url": "https://twitter.com/indiannavy"},
    {"handle": "@easaborterncomd", "name": "Eastern Command - Indian Army", "category": "defense", "url": "https://twitter.com/easterncomd"},
    {"handle": "@DefenceMinIndia", "name": "Ministry of Defence", "category": "government", "url": "https://twitter.com/DefenceMinIndia"},
    {"handle": "@MEAIndia", "name": "Ministry of External Affairs", "category": "government", "url": "https://twitter.com/MEAIndia"},
    {"handle": "@HMOIndia", "name": "Home Ministry", "category": "government", "url": "https://twitter.com/HMOIndia"},
    {"handle": "@PMOIndia", "name": "Prime Minister's Office", "category": "government", "url": "https://twitter.com/PMOIndia"},
    {"handle": "@BSF_India", "name": "Border Security Force", "category": "paramilitary", "url": "https://twitter.com/BSF_India"},
    {"handle": "@craborCRPF", "name": "CRPF", "category": "paramilitary", "url": "https://twitter.com/crpaborCRPF"},
    {"handle": "@official_dgar", "name": "Assam Rifles", "category": "paramilitary", "url": "https://twitter.com/official_dgar"},
    {"handle": "@ABORAITBP", "name": "ITBP", "category": "paramilitary", "url": "https://twitter.com/ITBP_official"},
    {"handle": "@SpsHanada", "name": "SPS Hanada - Defense Analyst", "category": "analyst", "url": "https://twitter.com/SpsHanada"},
]

THREAT_CATEGORIES = [
    "Insurgency", "Cross-border Movement", "Illegal Immigration",
    "Drug Trafficking", "Arms Smuggling", "Ethnic Conflicts",
    "Cyber Threats", "Strategic Infrastructure",
    "Political Developments", "Foreign Power Influence",
    "Military Operations", "Economic/Trade"
]
SEVERITY_LEVELS = ["low", "medium", "high", "critical"]
NER_STATES = ["Assam", "Meghalaya", "Mizoram", "Manipur", "Arunachal Pradesh", "Tripura"]
MONITORED_REGIONS = NER_STATES + ["Bangladesh", "Myanmar"]
BORDER_COUNTRIES = ["Bangladesh", "Myanmar"]


class IntelligenceItem(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    source: str
    source_url: str = ""
    published_at: str
    fetched_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    raw_content: str = ""
    ai_summary: str = ""
    why_it_matters: str = ""
    potential_impact: str = ""
    attention_level: str = "Routine Monitoring"
    state: str = ""
    threat_category: str = ""
    severity: str = "medium"
    is_cross_border: bool = False
    countries_involved: List[str] = []
    processed: bool = True
    tags: List[str] = []
    # New enhanced intelligence fields
    priority_score: int = 30
    confidence_score: int = 70
    threat_trajectory: str = "INDETERMINATE"
    regions: List[str] = []
    actors: List[str] = []
    special_flags: List[str] = []
    early_warning_signal: str = ""
    original_title: Optional[str] = None
    entities: Dict = Field(default_factory=lambda: {"persons": [], "organizations": [], "locations": []})
    # Sifter fields
    sifter_level: int = 1
    sifter_triggers: List[str] = []
    # Knowledge graph relationships
    relationships: List[Dict] = []
    # Vector embedding (stored separately, excluded from responses)
    embedding: Optional[List[float]] = Field(default=None, exclude=True)


# ============================================================
# WebSocket Connection Manager
# ============================================================
class ConnectionManager:
    """Manages WebSocket connections for real-time intelligence updates."""
    def __init__(self):
        self.active_connections: List[WebSocket] = []

    async def connect(self, websocket: WebSocket):
        await websocket.accept()
        self.active_connections.append(websocket)
        logger.info(f"WebSocket connected. Active: {len(self.active_connections)}")

    def disconnect(self, websocket: WebSocket):
        if websocket in self.active_connections:
            self.active_connections.remove(websocket)
        logger.info(f"WebSocket disconnected. Active: {len(self.active_connections)}")

    async def broadcast(self, message: dict):
        """Broadcast to all connected clients."""
        disconnected = []
        for conn in self.active_connections:
            try:
                await conn.send_json(message)
            except Exception:
                disconnected.append(conn)
        for conn in disconnected:
            self.disconnect(conn)

ws_manager = ConnectionManager()


class DailyBrief(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    date: str
    # NER Regional News
    key_developments: List[Dict] = []
    state_highlights: Dict[str, str] = {}
    cross_border_insights: str = ""
    analyst_summary: str = ""
    # National News Section
    national_news: List[Dict] = []
    # International News Section
    international_news: List[Dict] = []
    # Pattern Insights
    pattern_insights: List[Dict] = []
    # Uploaded Document Insights
    uploaded_insights: List[Dict] = []
    # Track included item IDs for cross-brief dedup
    included_item_ids: List[str] = []
    generated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())


class UploadedDocument(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    filename: str
    file_type: str
    uploaded_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    content_summary: str = ""
    extracted_text: str = ""
    ai_analysis: str = ""
    region: str = ""
    processed: bool = False


class TwitterFeed(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    handle: str
    account_name: str
    tweet_text: str
    tweet_url: str = ""
    posted_at: str
    fetched_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    category: str = "defense"
    is_relevant: bool = True


@api_router.get("/")
async def root():
    return {"message": "Rhino Drishti API - Intelligence Aggregation Platform"}


# ============================================================
# In-memory Cache for Dashboard Stats
# ============================================================
_stats_cache = {
    "data": None,
    "expires_at": None,
}
STATS_CACHE_TTL = 60  # seconds


def invalidate_stats_cache():
    _stats_cache["data"] = None
    _stats_cache["expires_at"] = None


@api_router.get("/dashboard/stats")
async def get_dashboard_stats():
    now = datetime.now(timezone.utc)
    if _stats_cache["data"] and _stats_cache["expires_at"] and now < _stats_cache["expires_at"]:
        return _stats_cache["data"]

    # Fetch retention setting
    settings = await db.app_settings.find_one({"key": "retention_days"}, {"_id": 0})
    retention_days = settings.get("value", 30) if settings else 30
    retention_cutoff = (now - timedelta(days=retention_days)).isoformat()
    base_filter = {"published_at": {"$gte": retention_cutoff}}

    total = await intelligence_col.count_documents(base_filter)
    critical = await intelligence_col.count_documents({**base_filter, "severity": "critical"})
    high = await intelligence_col.count_documents({**base_filter, "severity": "high"})
    medium = await intelligence_col.count_documents({**base_filter, "severity": "medium"})
    low = await intelligence_col.count_documents({**base_filter, "severity": "low"})

    state_dist = {}
    async for doc in intelligence_col.aggregate([
        {"$match": base_filter},
        {"$group": {"_id": "$state", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}}
    ]):
        if doc["_id"]:
            state_dist[doc["_id"]] = doc["count"]

    threat_dist = {}
    async for doc in intelligence_col.aggregate([
        {"$match": base_filter},
        {"$group": {"_id": "$threat_category", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}}
    ]):
        if doc["_id"]:
            threat_dist[doc["_id"]] = doc["count"]

    recent_critical = await intelligence_col.find(
        {**base_filter, "severity": {"$in": ["critical", "high"]}},
        {"_id": 0}
    ).sort("published_at", -1).limit(5).to_list(5)

    today = now.strftime("%Y-%m-%d")
    today_count = await intelligence_col.count_documents(
        {"published_at": {"$regex": f"^{today}"}}
    )

    trend_data = []
    async for doc in intelligence_col.aggregate([
        {"$match": base_filter},
        {"$group": {
            "_id": {"$substr": ["$published_at", 0, 10]},
            "count": {"$sum": 1},
            "critical": {"$sum": {"$cond": [{"$eq": ["$severity", "critical"]}, 1, 0]}},
            "high": {"$sum": {"$cond": [{"$eq": ["$severity", "high"]}, 1, 0]}}
        }},
        {"$sort": {"_id": 1}}
    ]):
        trend_data.append({"date": doc["_id"], "count": doc["count"], "critical": doc["critical"], "high": doc["high"]})

    result = {
        "total_items": total, "today_count": today_count,
        "critical_count": critical, "high_count": high,
        "medium_count": medium, "low_count": low,
        "state_distribution": state_dist, "threat_distribution": threat_dist,
        "recent_critical": recent_critical, "trend_7d": trend_data[-7:],
        "retention_days": retention_days,
    }

    _stats_cache["data"] = result
    _stats_cache["expires_at"] = now + timedelta(seconds=STATS_CACHE_TTL)
    return result


@api_router.get("/intelligence")
async def get_intelligence(
    state: Optional[str] = None,
    threat_type: Optional[str] = None,
    severity: Optional[str] = None,
    search: Optional[str] = None,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    is_cross_border: Optional[bool] = None,
    min_priority: Optional[int] = None,
    sort_by: Optional[str] = Query(None, description="Sort field: published_at, priority_score, severity"),
    sort_order: Optional[str] = Query("desc", description="Sort order: asc or desc"),
    page: int = Query(1, ge=1),
    limit: int = Query(20, ge=1, le=100),
    translate: bool = Query(True)
):
    query = {}

    # Apply retention window filter (unless explicit date_from is provided)
    if not date_from:
        settings = await db.app_settings.find_one({"key": "retention_days"}, {"_id": 0})
        retention_days = settings.get("value", 30) if settings else 30
        retention_cutoff = (datetime.now(timezone.utc) - timedelta(days=retention_days)).isoformat()
        query["published_at"] = {"$gte": retention_cutoff}

    if state:
        query["state"] = state
    if threat_type:
        query["threat_category"] = threat_type
    if severity:
        query["severity"] = severity
    if search:
        query["$or"] = [
            {"title": {"$regex": search, "$options": "i"}},
            {"ai_summary": {"$regex": search, "$options": "i"}},
            {"raw_content": {"$regex": search, "$options": "i"}}
        ]
    if date_from:
        query.setdefault("published_at", {})["$gte"] = date_from
    if date_to:
        query.setdefault("published_at", {})["$lte"] = date_to
    if is_cross_border is not None:
        query["is_cross_border"] = is_cross_border
    if min_priority is not None:
        query["priority_score"] = {"$gte": min_priority}

    # Sort logic
    sort_dir = -1 if sort_order == "desc" else 1
    sort_field = "published_at"
    if sort_by == "priority_score":
        sort_field = "priority_score"
    elif sort_by == "severity":
        sort_field = "severity"

    skip = (page - 1) * limit
    total = await intelligence_col.count_documents(query)
    items = await intelligence_col.find(query, {"_id": 0}).sort(sort_field, sort_dir).skip(skip).limit(limit).to_list(limit)

    # Translate non-English titles for display
    if translate:
        for item in items:
            if has_non_latin_chars(item.get("title", "")):
                item["title"] = await translate_to_english(item["title"])

    return {
        "items": items, "total": total, "page": page,
        "limit": limit, "pages": max((total + limit - 1) // limit, 0)
    }


@api_router.get("/intelligence/{item_id}")
async def get_intelligence_item(item_id: str):
    item = await intelligence_col.find_one({"id": item_id}, {"_id": 0})
    if not item:
        raise HTTPException(status_code=404, detail="Item not found")
    return item


@api_router.get("/alerts")
async def get_alerts():
    items = await intelligence_col.find(
        {"severity": {"$in": ["critical", "high"]}}, {"_id": 0}
    ).sort("published_at", -1).limit(30).to_list(30)
    return {"alerts": items, "count": len(items)}


@api_router.get("/alerts/unacknowledged")
async def get_unacknowledged_alerts():
    """Get critical/high alerts that have not been acknowledged."""
    items = await intelligence_col.find(
        {
            "severity": {"$in": ["critical", "high"]},
            "$or": [
                {"acknowledged": {"$exists": False}},
                {"acknowledged": False}
            ]
        },
        {"_id": 0}
    ).sort("published_at", -1).limit(50).to_list(50)
    return {"alerts": items, "count": len(items)}


@api_router.post("/intelligence/{item_id}/acknowledge")
async def acknowledge_alert(item_id: str):
    """Acknowledge a critical/high alert."""
    result = await intelligence_col.update_one(
        {"id": item_id},
        {"$set": {
            "acknowledged": True,
            "acknowledged_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Item not found")
    return {"message": "Alert acknowledged", "id": item_id}


@api_router.get("/patterns")
async def get_patterns():
    """Get detected intelligence patterns."""
    patterns = await patterns_col.find({}, {"_id": 0}).to_list(100)
    return {"patterns": patterns, "count": len(patterns)}


@api_router.post("/patterns/detect")
async def trigger_pattern_detection(background_tasks: BackgroundTasks):
    """Trigger pattern detection analysis."""
    from pattern_engine import detect_patterns
    background_tasks.add_task(detect_patterns, db)
    return {"message": "Pattern detection started"}


# ============================================================
# Settings Endpoints
# ============================================================
@api_router.get("/settings/retention")
async def get_retention_setting():
    """Get the current news retention window (in days)."""
    settings = await db.app_settings.find_one({"key": "retention_days"}, {"_id": 0})
    return {"retention_days": settings.get("value", 30) if settings else 30}


@api_router.put("/settings/retention")
async def set_retention_setting(body: dict):
    """Set the news retention window (in days). Valid range: 1-365."""
    days = body.get("retention_days", 30)
    if not isinstance(days, int) or days < 1 or days > 365:
        raise HTTPException(status_code=400, detail="retention_days must be integer 1-365")
    await db.app_settings.update_one(
        {"key": "retention_days"},
        {"$set": {"key": "retention_days", "value": days, "updated_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True
    )
    invalidate_stats_cache()
    return {"message": f"Retention window set to {days} days", "retention_days": days}


# ============================================================
# Semantic Search Endpoint
# ============================================================
@api_router.post("/intelligence/semantic-search")
async def intelligence_semantic_search(body: dict):
    """Semantic similarity search across intelligence items using vector embeddings."""
    query = body.get("query", "")
    limit = body.get("limit", 10)
    min_score = body.get("min_score", 0.3)
    
    if not query or len(query.strip()) < 3:
        raise HTTPException(status_code=400, detail="Query must be at least 3 characters")
    
    from embedding_service import semantic_search
    results = await semantic_search(db, query, limit=limit, min_score=min_score)
    
    return {"results": results, "count": len(results), "query": query}


@api_router.post("/embeddings/backfill")
async def trigger_embedding_backfill(background_tasks: BackgroundTasks):
    """Backfill vector embeddings for items that don't have them."""
    from embedding_service import backfill_embeddings
    background_tasks.add_task(backfill_embeddings, db, 50)
    return {"message": "Embedding backfill started (batch of 50)"}


# ============================================================
# On-Demand Custom PDF Brief
# ============================================================
@api_router.post("/intelligence/custom-brief")
async def generate_custom_brief(body: dict):
    """Generate a custom filtered PDF brief.
    
    Accepts filters: region, threat_type, severity, hours (last N hours), search
    """
    region = body.get("region")
    threat_type = body.get("threat_type")
    severity = body.get("severity")
    hours = body.get("hours", 48)
    search = body.get("search")
    title_override = body.get("title", "Custom Intelligence Brief")
    
    # Build query
    query = {"processed": True}
    cutoff = (datetime.now(timezone.utc) - timedelta(hours=hours)).isoformat()
    query["published_at"] = {"$gte": cutoff}
    
    if region:
        query["state"] = region
    if threat_type:
        query["threat_category"] = threat_type
    if severity:
        query["severity"] = severity
    if search:
        query["$or"] = [
            {"title": {"$regex": search, "$options": "i"}},
            {"ai_summary": {"$regex": search, "$options": "i"}}
        ]
    
    items = await intelligence_col.find(query, {"_id": 0}).sort(
        [("priority_score", -1), ("published_at", -1)]
    ).limit(50).to_list(50)
    
    if not items:
        raise HTTPException(status_code=404, detail="No items match the given filters")
    
    # Generate PDF
    pdf_bytes = _generate_custom_pdf(items, title_override, region, threat_type, hours)
    
    filename = f"custom_brief_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M')}.pdf"
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


def _generate_custom_pdf(items: list, title: str, region: str, threat: str, hours: int) -> bytes:
    """Generate a custom filtered PDF brief."""
    from fpdf import FPDF
    import pytz
    ist = pytz.timezone("Asia/Kolkata")
    now_ist = datetime.now(timezone.utc).astimezone(ist)
    
    class CustomPDF(FPDF):
        def header(self):
            # RESTRICTED header
            self.set_font('Helvetica', 'B', 8)
            self.set_text_color(200, 30, 30)
            self.cell(0, 4, 'RESTRICTED', align='C', new_x="LMARGIN", new_y="NEXT")
            self.ln(2)
        
        def footer(self):
            self.set_y(-15)
            self.set_font('Helvetica', 'I', 7)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, f'Page {self.page_no()} | RESTRICTED | Rhino Drishti Elite', align='C')
    
    pdf = CustomPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    
    # Title
    pdf.set_font('Helvetica', 'B', 16)
    pdf.set_text_color(34, 50, 30)
    clean_title = title.encode('latin-1', 'replace').decode('latin-1')
    pdf.cell(0, 10, clean_title, new_x="LMARGIN", new_y="NEXT")
    
    # Subtitle with filters
    pdf.set_font('Helvetica', '', 9)
    pdf.set_text_color(80, 80, 80)
    filter_desc = f"Generated: {now_ist.strftime('%d %b %Y %H:%M IST')} | Window: Last {hours}h"
    if region:
        filter_desc += f" | Region: {region}"
    if threat:
        filter_desc += f" | Threat: {threat}"
    pdf.cell(0, 5, filter_desc, new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 5, f"Total Items: {len(items)}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    
    # Items
    effective_w = pdf.w - pdf.l_margin - pdf.r_margin
    for i, item in enumerate(items, 1):
        sev = item.get('severity', 'low').upper()
        priority = item.get('priority_score', 0)
        trajectory = item.get('threat_trajectory', '')
        
        # Severity color
        if sev == 'CRITICAL':
            pdf.set_text_color(200, 30, 30)
        elif sev == 'HIGH':
            pdf.set_text_color(200, 100, 30)
        elif sev == 'MEDIUM':
            pdf.set_text_color(180, 160, 30)
        else:
            pdf.set_text_color(40, 120, 40)
        
        pdf.set_font('Helvetica', 'B', 10)
        item_title = item.get('title', 'Untitled')[:90]
        clean_item_title = item_title.encode('latin-1', 'replace').decode('latin-1')
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(effective_w, 6, f'{i}. [{sev}|P{priority}] {clean_item_title}')
        
        pdf.set_text_color(60, 60, 60)
        pdf.set_font('Helvetica', '', 8)
        meta = f"Source: {item.get('source', '')[:40]} | {item.get('state', '')} | {item.get('published_at', '')[:16]}"
        if trajectory and trajectory != 'INDETERMINATE':
            meta += f" | {trajectory}"
        clean_meta = meta.encode('latin-1', 'replace').decode('latin-1')
        pdf.set_x(pdf.l_margin)
        pdf.cell(effective_w, 4, clean_meta, new_x="LMARGIN", new_y="NEXT")
        
        summary = item.get('ai_summary', '')[:400]
        if summary:
            clean_summary = summary.encode('latin-1', 'replace').decode('latin-1')
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(effective_w, 4, clean_summary)
        
        why = item.get('why_it_matters', '')[:200]
        if why:
            pdf.set_font('Helvetica', 'I', 8)
            pdf.set_text_color(100, 80, 40)
            clean_why = why.encode('latin-1', 'replace').decode('latin-1')
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(effective_w, 4, f"Why: {clean_why}")
        
        pdf.ln(3)
    
    return pdf.output()


# ============================================================
# Web Scraping Trigger
# ============================================================
@api_router.post("/scrape-elite")
async def trigger_elite_scrape(background_tasks: BackgroundTasks):
    """Trigger elite web scraping of SATP, Ukhrul Times, Frontier Myanmar."""
    background_tasks.add_task(_run_elite_scrape)
    return {"message": "Elite scraping started"}


async def _run_elite_scrape():
    """Background task: scrape elite sources and ingest."""
    from web_scraper import scrape_all_targets
    articles = await scrape_all_targets()
    if articles:
        logger.info(f"Elite scrape: {len(articles)} articles found, merging into pipeline")
        # Dedup against existing
        for article in articles[:15]:
            existing = await intelligence_col.find_one(
                {"source_url": article["source_url"]}, {"_id": 1}
            )
            if not existing:
                article["processed"] = False
                await intelligence_col.insert_one(article)


@api_router.get("/daily-brief")
async def get_daily_brief(date: Optional[str] = None):
    if not date:
        date = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    brief = await briefs_col.find_one({"date": date}, {"_id": 0})
    if not brief:
        brief = await generate_brief_for_date(date)
    return brief


@api_router.post("/generate-brief")
async def generate_brief(background_tasks: BackgroundTasks):
    date = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    background_tasks.add_task(generate_brief_for_date, date)
    return {"message": "Brief generation started", "date": date}



@api_router.get("/daily-brief/pdf")
async def get_daily_brief_pdf(date: Optional[str] = None):
    """Generate and return a PDF of the daily intelligence brief with translated content"""
    if not date:
        date = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    brief = await briefs_col.find_one({"date": date}, {"_id": 0})
    if not brief:
        brief = await generate_brief_for_date(date)

    # Translate any non-English content in the brief for PDF
    brief = await translate_brief_for_pdf(brief)

    # Get stats for the PDF header
    total = await intelligence_col.count_documents({})
    critical = await intelligence_col.count_documents({"severity": "critical"})
    high = await intelligence_col.count_documents({"severity": "high"})

    # Query latest uploaded documents for this date directly from DB
    today_start = f"{date}T00:00:00"
    today_end = f"{date}T23:59:59"
    fresh_uploads = await uploads_col.find(
        {
            "processed": True,
            "uploaded_at": {"$gte": today_start, "$lte": today_end}
        },
        {"_id": 0}
    ).sort("uploaded_at", -1).to_list(20)

    pdf_bytes = generate_brief_pdf(brief, date, total, critical, high, fresh_uploads)

    import io
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=Rhino_Drishti_Brief_{date}.pdf"}
    )


async def translate_brief_for_pdf(brief: dict) -> dict:
    """Translate non-English content in brief to English for PDF rendering"""
    translated = dict(brief)
    
    # Translate key_developments
    if translated.get("key_developments"):
        new_devs = []
        for dev in translated["key_developments"]:
            if isinstance(dev, dict):
                if has_non_latin_chars(dev.get("title", "")):
                    dev["title"] = await translate_to_english(dev["title"])
                if has_non_latin_chars(dev.get("summary", "")):
                    dev["summary"] = await translate_to_english(dev["summary"])
                new_devs.append(dev)
            elif isinstance(dev, str) and has_non_latin_chars(dev):
                new_devs.append(await translate_to_english(dev))
            else:
                new_devs.append(dev)
        translated["key_developments"] = new_devs
    
    # Translate national_news
    if translated.get("national_news"):
        for news in translated["national_news"]:
            if isinstance(news, dict):
                if has_non_latin_chars(news.get("title", "")):
                    news["title"] = await translate_to_english(news["title"])
                if has_non_latin_chars(news.get("summary", "")):
                    news["summary"] = await translate_to_english(news["summary"])
    
    # Translate international_news
    if translated.get("international_news"):
        for news in translated["international_news"]:
            if isinstance(news, dict):
                if has_non_latin_chars(news.get("title", "")):
                    news["title"] = await translate_to_english(news["title"])
                if has_non_latin_chars(news.get("summary", "")):
                    news["summary"] = await translate_to_english(news["summary"])
    
    # Translate state_highlights
    if translated.get("state_highlights"):
        for state, highlight in translated["state_highlights"].items():
            if has_non_latin_chars(highlight):
                translated["state_highlights"][state] = await translate_to_english(highlight)
    
    # Translate analyst_summary and cross_border_insights
    if has_non_latin_chars(translated.get("analyst_summary", "")):
        translated["analyst_summary"] = await translate_to_english(translated["analyst_summary"])
    if has_non_latin_chars(translated.get("cross_border_insights", "")):
        translated["cross_border_insights"] = await translate_to_english(translated["cross_border_insights"])
    
    return translated


def has_non_latin_chars(text: str) -> bool:
    """Check if text contains non-Latin characters (Bengali, Hindi, Assamese, etc.)"""
    if not text:
        return False
    for char in text:
        code = ord(char)
        # Check for Bengali (U+0980-U+09FF), Devanagari (U+0900-U+097F), 
        # and other South Asian scripts
        if (0x0900 <= code <= 0x097F) or \
           (0x0980 <= code <= 0x09FF) or \
           (0x0A00 <= code <= 0x0A7F) or \
           (0x0A80 <= code <= 0x0AFF) or \
           (0x0B00 <= code <= 0x0B7F) or \
           (0x0B80 <= code <= 0x0BFF) or \
           (0x0C00 <= code <= 0x0C7F) or \
           (0x0C80 <= code <= 0x0CFF) or \
           (0x0D00 <= code <= 0x0D7F):
            return True
    return False


async def translate_to_english(text: str) -> str:
    """Translate non-English text to English using AI"""
    if not text or not has_non_latin_chars(text):
        return text
    
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY', '')
        
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"translate-{hash(text[:50])}",
            system_message="You are a translator. Translate the following text to English. Return ONLY the English translation, nothing else."
        ).with_model("anthropic", "claude-haiku-4-5-20251001")
        
        response = await chat.send_message(UserMessage(text=text[:1000]))
        return str(response).strip()
    except Exception as e:
        logger.error(f"Translation failed: {e}")
        # Return transliterated version as fallback
        return text.encode('ascii', 'ignore').decode('ascii') or "[Non-English content]"


def clean_for_pdf(text: str) -> str:
    """Clean text for PDF rendering - remove non-Latin characters"""
    if not text:
        return ""
    # Replace common Unicode characters with ASCII equivalents
    replacements = {
        '"': '"', '"': '"', ''': "'", ''': "'", 
        '–': '-', '—': '-', '…': '...', '•': '*',
        '\u200b': '', '\u200c': '', '\u200d': '',  # Zero-width chars
    }
    for orig, repl in replacements.items():
        text = text.replace(orig, repl)
    
    # If text has non-Latin chars, mark it for translation
    if has_non_latin_chars(text):
        # For PDF, we can't await, so we'll use a placeholder
        return "[Content requires translation - see original source]"
    
    return text.encode('latin-1', 'replace').decode('latin-1')


def generate_brief_pdf(brief: dict, date: str, total: int, critical: int, high: int, fresh_uploads: list = None) -> bytes:
    """Generate a professional PDF for the daily intelligence brief"""
    from fpdf import FPDF

    class BriefPDF(FPDF):
        def header(self):
            self.set_fill_color(30, 35, 25)
            self.rect(0, 0, 210, 40, 'F')
            self.set_font('Helvetica', 'B', 20)
            self.set_text_color(180, 220, 80)
            self.set_y(8)
            self.cell(0, 10, 'RHINO DRISHTI', align='C', new_x="LMARGIN", new_y="NEXT")
            self.set_font('Helvetica', '', 9)
            self.set_text_color(160, 170, 150)
            self.cell(0, 5, 'NER INTELLIGENCE PLATFORM  |  DAILY INTELLIGENCE BRIEF', align='C', new_x="LMARGIN", new_y="NEXT")
            self.set_font('Helvetica', '', 8)
            self.cell(0, 5, f'Classification: RESTRICTED  |  Date: {date}', align='C', new_x="LMARGIN", new_y="NEXT")
            self.ln(8)

        def footer(self):
            self.set_y(-15)
            self.set_font('Helvetica', 'I', 7)
            self.set_text_color(120, 120, 120)
            self.cell(0, 10, f'Rhino Drishti - Auto-generated Intelligence Brief | Page {self.page_no()}/{{nb}} | RESTRICTED', align='C')

        def section_title(self, title):
            self.set_font('Helvetica', 'B', 12)
            self.set_text_color(50, 60, 40)
            self.set_fill_color(230, 240, 220)
            self.cell(0, 8, f'  {title}', fill=True, new_x="LMARGIN", new_y="NEXT")
            self.ln(2)

        def body_text(self, text):
            self.set_font('Helvetica', '', 9)
            self.set_text_color(40, 40, 40)
            clean_text = text.encode('latin-1', 'replace').decode('latin-1')
            self.multi_cell(0, 5, clean_text)
            self.ln(2)
        
        def news_item_with_link(self, index, title, summary, source_url, timestamp=""):
            """Render a news item with embedded source link"""
            if self.get_y() > 260:
                self.add_page()
            
            self.set_font('Helvetica', 'B', 9)
            self.set_text_color(40, 60, 80)
            clean_title = title.encode('latin-1', 'replace').decode('latin-1')[:150]
            self.multi_cell(0, 5, f'{index}. {clean_title}', new_x="LMARGIN", new_y="NEXT")
            
            if summary:
                self.set_font('Helvetica', '', 8)
                self.set_text_color(60, 60, 60)
                clean_summary = summary.encode('latin-1', 'replace').decode('latin-1')[:400]
                if clean_summary:
                    self.multi_cell(0, 4, clean_summary, new_x="LMARGIN", new_y="NEXT")
            
            if source_url and len(source_url) > 5:
                self.set_font('Helvetica', 'I', 7)
                self.set_text_color(70, 100, 150)
                url_display = source_url[:70] + '...' if len(source_url) > 70 else source_url
                self.cell(0, 4, f'[Source: {url_display}]', new_x="LMARGIN", new_y="NEXT", link=source_url)
            
            if timestamp:
                self.set_font('Helvetica', 'I', 7)
                self.set_text_color(120, 120, 120)
                self.cell(0, 4, f'Time: {timestamp[:19]}', new_x="LMARGIN", new_y="NEXT")
            
            self.ln(2)
        
        def news_item_comprehensive(self, index, item):
            """Render a comprehensive news item with full analysis fields"""
            if self.get_y() > 240:
                self.add_page()
            
            title = item.get('title', '')
            summary = item.get('summary', '')
            source_url = item.get('source_url', '')
            severity = item.get('severity', '')
            state = item.get('state', '')
            priority = item.get('priority_score', 0)
            
            # Title with severity badge
            self.set_font('Helvetica', 'B', 9)
            self.set_text_color(40, 60, 80)
            sev_label = f' [{severity.upper()}]' if severity in ('critical', 'high') else ''
            clean_title = f'{index}. {title}{sev_label}'.encode('latin-1', 'replace').decode('latin-1')[:180]
            self.multi_cell(0, 5, clean_title, new_x="LMARGIN", new_y="NEXT")
            
            # Summary
            if summary:
                self.set_font('Helvetica', '', 8)
                self.set_text_color(60, 60, 60)
                self.multi_cell(0, 4, summary.encode('latin-1', 'replace').decode('latin-1')[:500], new_x="LMARGIN", new_y="NEXT")
            
            # Why it matters
            why = item.get('why_it_matters', '')
            if why:
                if self.get_y() > 270: self.add_page()
                self.set_font('Helvetica', 'B', 7)
                self.set_text_color(50, 120, 50)
                self.cell(30, 4, 'Why it matters: ', new_x="END")
                self.set_font('Helvetica', '', 7)
                self.set_text_color(80, 80, 80)
                self.multi_cell(0, 4, why.encode('latin-1', 'replace').decode('latin-1')[:300], new_x="LMARGIN", new_y="NEXT")
            
            # Potential impact
            impact = item.get('potential_impact', '')
            if impact:
                if self.get_y() > 270: self.add_page()
                self.set_font('Helvetica', 'B', 7)
                self.set_text_color(180, 120, 30)
                self.cell(30, 4, 'Potential impact: ', new_x="END")
                self.set_font('Helvetica', '', 7)
                self.set_text_color(80, 80, 80)
                self.multi_cell(0, 4, impact.encode('latin-1', 'replace').decode('latin-1')[:300], new_x="LMARGIN", new_y="NEXT")
            
            # Early warning
            warning = item.get('early_warning', '')
            if warning:
                if self.get_y() > 270: self.add_page()
                self.set_font('Helvetica', 'B', 7)
                self.set_text_color(200, 50, 50)
                self.cell(30, 4, 'EARLY WARNING: ', new_x="END")
                self.set_font('Helvetica', '', 7)
                self.set_text_color(150, 50, 50)
                self.multi_cell(0, 4, warning.encode('latin-1', 'replace').decode('latin-1')[:300], new_x="LMARGIN", new_y="NEXT")
            
            # Special flags
            flags = item.get('special_flags', [])
            if flags and isinstance(flags, list) and len(flags) > 0:
                self.set_font('Helvetica', 'I', 7)
                self.set_text_color(140, 100, 40)
                flags_text = 'Flags: ' + ' | '.join(str(f) for f in flags[:5])
                self.cell(0, 4, flags_text.encode('latin-1', 'replace').decode('latin-1')[:150], new_x="LMARGIN", new_y="NEXT")
            
            # Actors
            actors = item.get('actors', '')
            if actors:
                self.set_font('Helvetica', 'I', 7)
                self.set_text_color(100, 100, 120)
                self.cell(0, 4, f'Actors: {str(actors)[:120]}'.encode('latin-1', 'replace').decode('latin-1'), new_x="LMARGIN", new_y="NEXT")
            
            # Source link and metadata
            meta_parts = []
            if state: meta_parts.append(f'Region: {state}')
            if priority: meta_parts.append(f'Priority: {priority}')
            if source_url:
                url_display = source_url[:60] + '...' if len(source_url) > 60 else source_url
                meta_parts.append(f'Source: {url_display}')
            
            if meta_parts:
                self.set_font('Helvetica', 'I', 7)
                self.set_text_color(100, 100, 100)
                meta_text = ' | '.join(meta_parts)
                if source_url:
                    self.cell(0, 4, meta_text.encode('latin-1', 'replace').decode('latin-1')[:180], new_x="LMARGIN", new_y="NEXT", link=source_url)
                else:
                    self.cell(0, 4, meta_text.encode('latin-1', 'replace').decode('latin-1')[:180], new_x="LMARGIN", new_y="NEXT")
            
            self.ln(3)

    pdf = BriefPDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Situation Overview box
    pdf.set_fill_color(245, 245, 240)
    pdf.set_draw_color(180, 190, 170)
    pdf.rect(10, pdf.get_y(), 190, 18, 'FD')
    y_start = pdf.get_y() + 3
    pdf.set_xy(15, y_start)
    pdf.set_font('Helvetica', 'B', 9)
    pdf.set_text_color(40, 40, 40)
    pdf.cell(60, 5, f'Total Intelligence Items: {total}')
    pdf.cell(60, 5, f'Critical Alerts: {critical}')
    pdf.cell(60, 5, f'High Priority: {high}')
    pdf.set_xy(15, y_start + 7)
    pdf.set_font('Helvetica', '', 8)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 5, 'Covering: Assam, Meghalaya, Mizoram, Manipur, Arunachal Pradesh, Tripura')
    pdf.ln(15)

    # ========== NER REGIONAL SECTION (ONLY SECTION) ==========
    pdf.section_title('NORTHEAST REGION - KEY DEVELOPMENTS')
    
    # Filter key developments to NER-only items
    NER_STATES_PDF = ["Assam", "Meghalaya", "Mizoram", "Manipur", "Arunachal Pradesh", "Tripura", "Multiple", ""]
    developments = brief.get('key_developments', [])
    ner_developments = []
    for dev in developments:
        if isinstance(dev, dict):
            state = dev.get('state', '')
            # Include NER states, or items without a state (likely NER-related from the query)
            if state in NER_STATES_PDF or not state:
                ner_developments.append(dev)
        else:
            ner_developments.append(dev)
    
    if ner_developments:
        for i, dev in enumerate(ner_developments, 1):
            if isinstance(dev, dict):
                pdf.news_item_comprehensive(i, dev)
            else:
                pdf.set_font('Helvetica', '', 9)
                pdf.set_text_color(40, 40, 40)
                clean_dev = str(dev).encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 5, f'{i:02d}. {clean_dev}')
                pdf.ln(1)
    else:
        pdf.body_text('No key developments recorded for this period.')
    pdf.ln(2)

    # ========== UPLOADED DOCUMENT INSIGHTS (same date only, NER-focused) ==========
    ner_keywords = ['assam', 'meghalaya', 'mizoram', 'manipur', 'arunachal', 'tripura', 'nagaland', 
                    'northeast', 'ner', 'nscn', 'ulfa', 'pla', 'rpf', 'myanmar', 'border',
                    'insurgency', 'militant', 'security force', 'army', 'bsf', 'crpf', 'assam rifles']
    same_date_docs = []
    for doc in (fresh_uploads or []):
        analysis_text = (str(doc.get('ai_analysis', '')) + str(doc.get('extracted_text', '')) + str(doc.get('filename', ''))).lower()
        if any(kw in analysis_text for kw in ner_keywords):
            same_date_docs.append(doc)
    
    if same_date_docs:
        pdf.section_title('UPLOADED DOCUMENT INSIGHTS')
        for i, doc in enumerate(same_date_docs[:10], 1):
            pdf.set_font('Helvetica', 'B', 9)
            pdf.set_text_color(80, 60, 40)
            filename = doc.get('filename', 'Unknown Document')
            pdf.cell(0, 5, f'{i}. {filename}', new_x="LMARGIN", new_y="NEXT")
            
            pdf.set_font('Helvetica', '', 8)
            pdf.set_text_color(60, 60, 60)
            analysis = doc.get('ai_analysis', doc.get('content_summary', ''))[:500]
            clean_analysis = analysis.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 4, clean_analysis)
            pdf.ln(2)
    
    # ========== PATTERN INSIGHTS (Escalation Warnings) ==========
    pattern_insights = brief.get('pattern_insights', [])
    if pattern_insights:
        pdf.section_title('PATTERN DETECTION - ESCALATION WARNINGS')
        for i, p in enumerate(pattern_insights[:10], 1):
            risk = p.get('escalation_risk', 'LOW')
            region = p.get('region', 'Unknown')
            detail = p.get('detail', p.get('pattern_type', ''))
            events = p.get('event_count', 0)
            avg_pri = p.get('avg_priority_score', 0)
            window = p.get('window_days', 7)
            
            # Risk color
            if risk == 'CRITICAL':
                pdf.set_text_color(200, 30, 30)
            elif risk == 'HIGH':
                pdf.set_text_color(200, 100, 30)
            elif risk == 'MODERATE':
                pdf.set_text_color(180, 160, 30)
            else:
                pdf.set_text_color(40, 120, 40)
            
            pdf.set_font('Helvetica', 'B', 9)
            header = f'{i}. [{risk}] {region} - {detail}'
            clean_header = header.encode('latin-1', 'replace').decode('latin-1')
            pdf.cell(0, 5, clean_header, new_x="LMARGIN", new_y="NEXT")
            
            pdf.set_font('Helvetica', '', 8)
            pdf.set_text_color(80, 80, 80)
            stats_line = f'   {events} events in {window} days | Avg Priority: {avg_pri}'
            pdf.cell(0, 4, stats_line, new_x="LMARGIN", new_y="NEXT")
            
            # Sample titles
            samples = p.get('sample_titles', [])
            if samples:
                pdf.set_font('Helvetica', 'I', 7)
                pdf.set_text_color(100, 100, 100)
                for title in samples[:2]:
                    clean_t = str(title).encode('latin-1', 'replace').decode('latin-1')[:120]
                    pdf.cell(0, 4, f'     - {clean_t}', new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
    
    # ========== ANALYST SUMMARY ==========
    pdf.add_page()
    pdf.section_title('ANALYST ASSESSMENT')
    summary = brief.get('analyst_summary', 'No analyst summary available.')
    pdf.body_text(summary)
    pdf.ln(2)

    # Classification footer
    pdf.ln(5)
    pdf.set_font('Helvetica', 'B', 8)
    pdf.set_text_color(150, 50, 50)
    pdf.cell(0, 5, 'DISTRIBUTION: RESTRICTED | FOR AUTHORIZED PERSONNEL ONLY', align='C')

    return pdf.output()


@api_router.get("/weekly-trends")
async def get_weekly_trends():
    daily_severity = {}
    async for doc in intelligence_col.aggregate([
        {"$group": {
            "_id": {"date": {"$substr": ["$published_at", 0, 10]}, "severity": "$severity"},
            "count": {"$sum": 1}
        }},
        {"$sort": {"_id.date": 1}}
    ]):
        date = doc["_id"]["date"]
        sev = doc["_id"]["severity"]
        if date not in daily_severity:
            daily_severity[date] = {"date": date, "critical": 0, "high": 0, "medium": 0, "low": 0, "total": 0}
        daily_severity[date][sev] = doc["count"]
        daily_severity[date]["total"] += doc["count"]

    category_stats = []
    async for doc in intelligence_col.aggregate([
        {"$group": {
            "_id": "$threat_category", "count": {"$sum": 1},
            "critical": {"$sum": {"$cond": [{"$eq": ["$severity", "critical"]}, 1, 0]}},
            "high": {"$sum": {"$cond": [{"$eq": ["$severity", "high"]}, 1, 0]}}
        }},
        {"$sort": {"count": -1}}
    ]):
        if doc["_id"]:
            category_stats.append({"category": doc["_id"], "count": doc["count"], "critical": doc["critical"], "high": doc["high"]})

    state_stats = []
    async for doc in intelligence_col.aggregate([
        {"$group": {
            "_id": "$state", "count": {"$sum": 1},
            "critical": {"$sum": {"$cond": [{"$eq": ["$severity", "critical"]}, 1, 0]}},
            "high": {"$sum": {"$cond": [{"$eq": ["$severity", "high"]}, 1, 0]}}
        }},
        {"$sort": {"count": -1}}
    ]):
        if doc["_id"]:
            state_stats.append({"state": doc["_id"], "count": doc["count"], "critical": doc["critical"], "high": doc["high"]})

    return {
        "daily_severity": sorted(daily_severity.values(), key=lambda x: x["date"])[-14:],
        "category_stats": category_stats,
        "state_stats": state_stats
    }


@api_router.get("/sources")
async def get_sources():
    sources = await sources_col.find({}, {"_id": 0}).to_list(100)
    return {"sources": sources}


@api_router.get("/twitter-accounts")
async def get_twitter_accounts():
    """Get list of Twitter/X accounts being monitored"""
    return {"accounts": TWITTER_ACCOUNTS_TO_MONITOR}


@api_router.get("/twitter-feeds")
async def get_twitter_feeds(limit: int = Query(50, ge=1, le=200)):
    """Get recent Twitter/X feeds from monitored accounts"""
    feeds = await tweets_col.find({}, {"_id": 0}).sort("posted_at", -1).limit(limit).to_list(limit)
    return {"feeds": feeds, "count": len(feeds)}


@api_router.get("/uploaded-documents")
async def get_uploaded_documents():
    """Get list of uploaded documents"""
    docs = await uploads_col.find({}, {"_id": 0}).sort("uploaded_at", -1).to_list(100)
    return {"documents": docs, "count": len(docs)}


@api_router.post("/upload-document")
async def upload_document(file: UploadFile = File(...), background_tasks: BackgroundTasks = None):
    """Upload a PDF, Word, or Excel document for intelligence analysis"""
    allowed_types = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/msword": "doc",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
        "application/vnd.ms-excel": "xls",
        "text/plain": "txt"
    }
    
    content_type = file.content_type
    if content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="File type not supported. Allowed: PDF, Word, Excel, TXT")
    
    file_type = allowed_types[content_type]
    file_content = await file.read()
    
    # Extract text from document
    extracted_text = ""
    try:
        if file_type == "pdf":
            from PyPDF2 import PdfReader
            pdf_reader = PdfReader(io.BytesIO(file_content))
            for page in pdf_reader.pages:
                extracted_text += page.extract_text() or ""
        elif file_type in ["docx", "doc"]:
            from docx import Document
            doc = Document(io.BytesIO(file_content))
            extracted_text = "\n".join([para.text for para in doc.paragraphs])
        elif file_type in ["xlsx", "xls"]:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(file_content))
            for sheet in wb:
                for row in sheet.iter_rows(values_only=True):
                    extracted_text += " | ".join([str(cell) for cell in row if cell]) + "\n"
        elif file_type == "txt":
            extracted_text = file_content.decode('utf-8', errors='ignore')
    except Exception as e:
        logger.error(f"Error extracting text from {file.filename}: {e}")
        extracted_text = f"Error extracting text: {str(e)}"
    
    # Create document record
    doc_record = {
        "id": str(uuid.uuid4()),
        "filename": file.filename,
        "file_type": file_type,
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
        "content_summary": extracted_text[:500] + "..." if len(extracted_text) > 500 else extracted_text,
        "extracted_text": extracted_text[:10000],  # Limit to 10k chars
        "ai_analysis": "",
        "region": "",
        "processed": False
    }
    
    await uploads_col.insert_one(doc_record)
    
    # Trigger AI analysis in background
    if background_tasks:
        background_tasks.add_task(analyze_uploaded_document, doc_record["id"])
    
    return {
        "message": "Document uploaded successfully",
        "document_id": doc_record["id"],
        "filename": file.filename,
        "extracted_chars": len(extracted_text)
    }


@api_router.delete("/uploaded-documents/{doc_id}")
async def delete_uploaded_document(doc_id: str):
    """Delete an uploaded document"""
    result = await uploads_col.delete_one({"id": doc_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"message": "Document deleted successfully"}


@api_router.post("/fetch-news")
async def trigger_fetch(background_tasks: BackgroundTasks):
    background_tasks.add_task(fetch_and_process_news)
    return {"message": "News fetch triggered"}


@api_router.post("/bulk-scrape")
async def trigger_bulk_scrape(background_tasks: BackgroundTasks):
    """Bulk scrape ALL articles from RSS feeds without AI processing.
    Articles are stored as unprocessed and will be AI-analyzed gradually."""
    background_tasks.add_task(bulk_scrape_all_feeds)
    return {"message": "Bulk scrape triggered - articles will be stored for gradual AI processing"}


@api_router.get("/scan-status")
async def get_scan_status():
    """Get real-time RSS scan status"""
    return scan_status


@api_router.post("/analyze-news")
async def trigger_analysis(background_tasks: BackgroundTasks):
    background_tasks.add_task(analyze_unprocessed_items)
    return {"message": "Analysis triggered"}


@api_router.get("/pipeline/status")
async def pipeline_status():
    """Show processing pipeline health with rate limit management info"""
    total = await intelligence_col.count_documents({})
    processed = await intelligence_col.count_documents({"processed": True})
    unprocessed = await intelligence_col.count_documents({"processed": False})
    sources = await sources_col.count_documents({})
    
    # Get recent processing stats (last 24 hours)
    yesterday = (datetime.now(timezone.utc) - timedelta(hours=24)).isoformat()
    recent_processed = await intelligence_col.count_documents({
        "processed": True,
        "fetched_at": {"$gte": yesterday}
    })
    
    return {
        "total_items": total,
        "ai_processed": processed,
        "pending_retry": unprocessed,
        "processing_rate": f"{(processed / total * 100):.1f}%" if total > 0 else "N/A",
        "recent_24h_processed": recent_processed,
        "rss_sources": sources,
        "rate_limit_config": {
            "max_articles_per_cycle": 25,
            "batch_size": 3,
            "batch_pause_seconds": 5,
            "inter_article_delay_seconds": 1.5,
            "max_retry_per_cycle": 15
        },
        "filter_stats": {
            "last_filtered_out": scan_status.get("filtered_out", 0),
            "last_translated": scan_status.get("translated", 0),
        },
        "scheduler": "grassroots/60min, standard/30min, established/12hr, retry/15min, brief/0600 IST, embeddings/6hr"
    }



async def generate_brief_for_date(date: str):
    """Generate comprehensive daily brief.
    
    Time Window Logic:
    - Auto 0600h brief: covers 0600h IST previous day → 0600h IST today
    - Manual regeneration after 0600h: covers from the LATEST brief generated 
      on the PREVIOUS calendar date → current time
    - No item that appeared in a previous brief will be repeated.
    """
    
    # Define NER states (India's Northeast Region)
    NER_STATES = ["Assam", "Meghalaya", "Mizoram", "Manipur", "Arunachal Pradesh", "Tripura"]
    
    # ========== TIME WINDOW CALCULATION ==========
    from datetime import timedelta
    import pytz
    ist = pytz.timezone("Asia/Kolkata")
    now_utc = datetime.now(timezone.utc)
    now_ist = now_utc.astimezone(ist)
    
    # Find the LATEST brief from the PREVIOUS calendar date
    previous_date = (now_ist - timedelta(days=1)).strftime("%Y-%m-%d")
    
    # Get all briefs from previous calendar date, sorted by generation time (latest first)
    prev_day_brief = await briefs_col.find_one(
        {"date": previous_date},
        {"_id": 0, "generated_at": 1, "included_item_ids": 1, "date": 1},
        sort=[("generated_at", -1)]
    )
    
    # Also check any earlier briefs (e.g. if previous day had no brief)
    any_previous_brief = await briefs_col.find_one(
        {"date": {"$lt": date}},
        {"_id": 0, "generated_at": 1, "included_item_ids": 1, "date": 1},
        sort=[("date", -1)]
    )
    
    # Collect item IDs from previous briefs to avoid repeats
    previous_item_ids = set()
    
    # Determine time cutoff based on scenario
    today_0600_ist = now_ist.replace(hour=6, minute=0, second=0, microsecond=0)
    
    if prev_day_brief and prev_day_brief.get("generated_at"):
        # SCENARIO: Previous day has a brief — use its generation time as cutoff
        cutoff_utc = prev_day_brief["generated_at"]
        if prev_day_brief.get("included_item_ids"):
            previous_item_ids.update(prev_day_brief["included_item_ids"])
        logger.info(f"Brief window: prev day brief ({previous_date}) generated at {cutoff_utc} → now")
    elif any_previous_brief and any_previous_brief.get("generated_at"):
        # SCENARIO: No brief yesterday but older briefs exist
        cutoff_utc = any_previous_brief["generated_at"]
        if any_previous_brief.get("included_item_ids"):
            previous_item_ids.update(any_previous_brief["included_item_ids"])
        logger.info(f"Brief window: last brief ({any_previous_brief.get('date')}) generated at {cutoff_utc} → now")
    else:
        # SCENARIO: First-ever brief — default to 0600h IST previous day
        cutoff_ist = today_0600_ist - timedelta(days=1)
        cutoff_utc = cutoff_ist.astimezone(timezone.utc).isoformat()
        logger.info(f"Brief window (first brief): {cutoff_utc} → now")
    
    logger.info(f"Excluding {len(previous_item_ids)} items from previous briefs")
    
    # ========== TITLE SIMILARITY DEDUP HELPER ==========
    def normalize_title(title):
        """Normalize title for similarity comparison"""
        import re
        t = (title or "").lower().strip()
        t = re.sub(r'[^a-z0-9\s]', '', t)
        t = re.sub(r'\s+', ' ', t).strip()
        return t
    
    def extract_key_entities(title):
        """Extract key entities (names, places, orgs) for event matching"""
        import re
        t = (title or "").lower()
        # Known NER entities
        entities = set()
        known_orgs = ['ulfa', 'nscn', 'rpf', 'pla', 'hnlc', 'gnla', 'kla', 'mnf', 'unlf', 'prepak', 'knf', 'arsa', 'tnla', 'mndaa', 'assam rifles', 'bsf', 'crpf', 'army']
        known_places = ['manipur', 'assam', 'meghalaya', 'mizoram', 'tripura', 'arunachal', 'nagaland', 'tinsukia', 'changlang', 'tamenglong', 'imphal', 'dimapur', 'guwahati', 'silchar', 'agartala', 'shillong', 'aizawl', 'itanagar', 'kohima', 'myanmar', 'bangladesh', 'dhaka', 'chittagong', 'cox']
        known_events = ['rpg', 'grenade', 'gunfire', 'gunfight', 'bomb', 'blast', 'attack', 'ambush', 'shootout', 'firing', 'seized', 'arrested', 'killed', 'injured', 'rally', 'protest', 'blockade', 'bandh']
        
        for org in known_orgs:
            if org in t:
                entities.add(org)
        for place in known_places:
            if place in t:
                entities.add(place)
        for event in known_events:
            if event in t:
                entities.add(event)
        # Extract numbers (casualties, etc.)
        numbers = re.findall(r'\b(\d+)\s*(?:killed|injured|arrested|seized|dead)\b', t)
        for n in numbers:
            entities.add(f"count_{n}")
        return entities
    
    def is_duplicate_title(new_title, seen_titles, seen_entities_list, threshold=0.55):
        """Check if title is similar to any already seen title using word overlap AND entity matching"""
        norm_new = normalize_title(new_title)
        if not norm_new or len(norm_new) < 10:
            return False
        new_words = set(norm_new.split())
        new_entities = extract_key_entities(new_title)
        
        for i, seen in enumerate(seen_titles):
            seen_words = set(seen.split())
            if not seen_words:
                continue
            # Word overlap check
            overlap = len(new_words & seen_words)
            total = max(len(new_words), len(seen_words))
            word_sim = overlap / total if total > 0 else 0
            
            # Entity overlap check (stricter - same event if key entities match)
            seen_ents = seen_entities_list[i] if i < len(seen_entities_list) else set()
            if new_entities and seen_ents:
                ent_overlap = len(new_entities & seen_ents)
                ent_total = max(len(new_entities), len(seen_ents))
                ent_sim = ent_overlap / ent_total if ent_total > 0 else 0
                
                # If 3+ entities match (e.g. ULFA + Tinsukia + attack), it's the same event
                if ent_overlap >= 3:
                    return True
                # If 2 entities match AND word similarity is moderate
                if ent_overlap >= 2 and word_sim >= 0.35:
                    return True
            
            # Standard word similarity
            if word_sim >= threshold:
                return True
        return False
    
    # ========== 1. GET ALL CRITICAL/HIGH ITEMS since last brief ==========
    # Base query: processed items in time window, excluding items from previous briefs
    base_filter = {"processed": True, "published_at": {"$gte": cutoff_utc}}
    if previous_item_ids:
        base_filter["id"] = {"$nin": list(previous_item_ids)}
    
    critical_high_items = await intelligence_col.find(
        {
            **base_filter,
            "$or": [
                {"severity": {"$in": ["critical", "high"]}},
                {"priority_score": {"$gte": 60}}
            ]
        },
        {"_id": 0}
    ).sort([("priority_score", -1), ("published_at", -1)]).to_list(100)
    
    # Fallback: if too few items in time window, expand but STILL exclude previous brief items
    if len(critical_high_items) < 5:
        logger.info(f"Brief: Only {len(critical_high_items)} items in time window, expanding to recent critical/high")
        fallback_filter = {"processed": True}
        if previous_item_ids:
            fallback_filter["id"] = {"$nin": list(previous_item_ids)}
        fallback_items = await intelligence_col.find(
            {
                **fallback_filter,
                "$or": [
                    {"severity": {"$in": ["critical", "high"]}},
                    {"priority_score": {"$gte": 60}}
                ]
            },
            {"_id": 0}
        ).sort([("priority_score", -1), ("published_at", -1)]).limit(30).to_list(30)
        # Merge without duplicates
        existing_urls = set(i.get("source_url") for i in critical_high_items)
        for item in fallback_items:
            if item.get("source_url") not in existing_urls:
                critical_high_items.append(item)
                existing_urls.add(item.get("source_url"))
    
    # Dedup critical/high items by title similarity
    seen_titles = []
    seen_entities = []
    deduped_critical = []
    for item in critical_high_items:
        title = item.get("title", "")
        if not is_duplicate_title(title, seen_titles, seen_entities):
            deduped_critical.append(item)
            seen_titles.append(normalize_title(title))
            seen_entities.append(extract_key_entities(title))
    
    logger.info(f"Brief: {len(critical_high_items)} critical/high items found, {len(deduped_critical)} after title dedup")
    
    # ========== 2. GET NER REGIONAL ITEMS with time window ==========
    ner_query = {
        "processed": True,
        "published_at": {"$gte": cutoff_utc},
        "state": {"$in": NER_STATES + ["Multiple"]},
        "$or": [
            {"priority_score": {"$gte": 30}},
            {"tags": {"$exists": True, "$ne": []}},
            {"severity": {"$in": ["critical", "high", "medium"]}}
        ]
    }
    if previous_item_ids:
        ner_query["id"] = {"$nin": list(previous_item_ids)}
    
    ner_items = await intelligence_col.find(
        ner_query, {"_id": 0}
    ).sort([("priority_score", -1), ("published_at", -1)]).limit(80).to_list(80)
    
    # Fallback: if too few NER items in time window
    if len(ner_items) < 5:
        logger.info(f"Brief: Only {len(ner_items)} NER items in window, expanding")
        fallback_ner_query = {
            "processed": True,
            "state": {"$in": NER_STATES + ["Multiple"]},
            "severity": {"$in": ["critical", "high", "medium"]}
        }
        if previous_item_ids:
            fallback_ner_query["id"] = {"$nin": list(previous_item_ids)}
        fallback_ner = await intelligence_col.find(
            fallback_ner_query, {"_id": 0}
        ).sort([("priority_score", -1), ("published_at", -1)]).limit(40).to_list(40)
        existing_urls = set(i.get("source_url") for i in ner_items)
        for item in fallback_ner:
            if item.get("source_url") not in existing_urls:
                ner_items.append(item)
                existing_urls.add(item.get("source_url"))
    
    # Dedup NER items and diversify sources
    seen_sources = {}
    diverse_ner_items = []
    for item in ner_items:
        title = item.get("title", "")
        source = item.get("source", "Unknown")
        if is_duplicate_title(title, seen_titles, seen_entities):
            continue
        if source not in seen_sources:
            seen_sources[source] = 0
        if seen_sources[source] < 4:
            diverse_ner_items.append(item)
            seen_sources[source] += 1
            seen_titles.append(normalize_title(title))
            seen_entities.append(extract_key_entities(title))
    
    logger.info(f"Brief: {len(ner_items)} NER items, {len(diverse_ner_items)} after dedup from {len(seen_sources)} sources")
    
    # ========== 3. GET NATIONAL NEWS with time window ==========
    national_query = {
        "processed": True,
        "published_at": {"$gte": cutoff_utc},
        "source": {"$in": ["The Hindu - National", "NDTV India News", "News18 India", "Times of India", "PIB Press Releases", "PIB Defence", "MHA India"]},
        "state": {"$nin": NER_STATES + ["Bangladesh", "Myanmar", "Multiple"]}
    }
    if previous_item_ids:
        national_query["id"] = {"$nin": list(previous_item_ids)}
    national_items = await intelligence_col.find(
        national_query, {"_id": 0}
    ).sort([("priority_score", -1), ("published_at", -1)]).limit(30).to_list(30)
    
    military_national = [
        item for item in national_items
        if item.get("priority_score", 0) >= 25 or 
           item.get("severity") in ["critical", "high", "medium"] or
           any(tag in str(item.get("tags", [])).lower() for tag in ["military", "security", "cross-border", "insurgency", "foreign", "infrastructure", "defence", "border"])
    ]
    
    logger.info(f"Brief: {len(national_items)} national items, {len(military_national)} military-relevant")
    
    # ========== 4. GET INTERNATIONAL NEWS with time window ==========
    intl_query = {
        "processed": True,
        "published_at": {"$gte": cutoff_utc},
        "$or": [
            {"state": {"$in": ["Bangladesh", "Myanmar"]}},
            {"countries_involved": {"$in": ["China", "Pakistan"]}},
            {"priority_score": {"$gte": 35}},
            {"severity": {"$in": ["critical", "high"]}},
            {"tags": {"$in": [
                "Military Movement", "Cross-border Movement", "Insurgency / Militancy",
                "Foreign Influence (China/Pakistan/USA)", "Border Security", "Arms Smuggling",
                "Drug Trafficking", "Illegal Immigration", "Bangladesh Internal Dynamics",
                "Myanmar Instability", "Infrastructure / Logistics"
            ]}}
        ],
        "state": {"$nin": NER_STATES + ["Multiple", "India", ""]}
    }
    if previous_item_ids:
        intl_query["id"] = {"$nin": list(previous_item_ids)}
    international_items = await intelligence_col.find(
        intl_query, {"_id": 0}
    ).sort([("priority_score", -1), ("published_at", -1)]).limit(40).to_list(40)
    
    EXCLUDE_KEYWORDS = [
        'cricket', 'football', 'sports', 'match', 'tournament', 'celebrity', 'entertainment',
        'movie', 'film', 'music', 'concert', 'festival', 'recipe', 'fashion', 'lifestyle',
        'wedding', 'divorce', 'sparrow', 'bird', 'animal', 'zoo', 'weather forecast',
        'horoscope', 'lottery', 'quiz', 'game show', 'reality show', 'bollywood', 'tollywood',
        'chelsea', 'goalkeeper', 'striker', 'midfielder', 'coach', 'player'
    ]
    
    def is_strategic_news(item):
        title = (item.get("title", "") or "").lower()
        summary = (item.get("ai_summary", "") or "").lower()
        content = title + " " + summary
        for kw in EXCLUDE_KEYWORDS:
            if kw in content:
                return False
        if item.get("priority_score", 0) >= 35:
            return True
        tags = item.get("tags", [])
        security_tags = ["Military", "Border", "Insurgency", "Cross-border", "Arms", "Drug", "Security", "Foreign"]
        for tag in tags:
            for st in security_tags:
                if st.lower() in tag.lower():
                    return True
        return False
    
    strategic_intl_items = [item for item in international_items if is_strategic_news(item)]
    
    # Dedup international items
    seen_intl_sources = {}
    diverse_intl_items = []
    for item in strategic_intl_items:
        title = item.get("title", "")
        source = item.get("source", "Unknown")
        if is_duplicate_title(title, seen_titles, seen_entities):
            continue
        if source not in seen_intl_sources:
            seen_intl_sources[source] = 0
        if seen_intl_sources[source] < 3:
            diverse_intl_items.append(item)
            seen_intl_sources[source] += 1
            seen_titles.append(normalize_title(title))
            seen_entities.append(extract_key_entities(title))
    
    logger.info(f"Brief: {len(international_items)} intl items, {len(strategic_intl_items)} strategic, {len(diverse_intl_items)} deduplicated")
    
    # ========== 5. GET PATTERN INSIGHTS ==========
    from pattern_engine import detect_patterns
    try:
        detected_patterns = await detect_patterns(db)
    except Exception as e:
        logger.warning(f"Pattern detection failed during brief gen: {e}")
        detected_patterns = await patterns_col.find({}, {"_id": 0}).to_list(50)
    
    # ========== 6. GET UPLOADED DOCUMENT INSIGHTS ==========
    uploaded_docs = await uploads_col.find({"processed": True}, {"_id": 0}).sort("uploaded_at", -1).limit(10).to_list(10)
    
    # ========== 7. GENERATE AI BRIEF ==========
    items_for_ai = deduped_critical + diverse_ner_items[:20]
    
    try:
        from ai_pipeline import generate_daily_brief_ai
        brief_data = await generate_daily_brief_ai(items_for_ai, date)
    except Exception as e:
        logger.error(f"AI brief generation failed: {e}")
        brief_data = generate_manual_brief(items_for_ai, date)
    
    # ========== 8. BUILD COMPREHENSIVE KEY DEVELOPMENTS ==========
    # Helper to build a comprehensive item dict with all analysis fields
    def build_brief_item(item):
        """Build a comprehensive brief item including analysis and pattern detection"""
        result = {
            "title": item.get("title", ""),
            "summary": item.get("ai_summary", ""),
            "source_url": item.get("source_url", ""),
            "timestamp": item.get("published_at", ""),
            "severity": item.get("severity", "medium"),
            "priority_score": item.get("priority_score", 0),
            "state": item.get("state", ""),
            "source": item.get("source", ""),
        }
        # Include analysis fields
        if item.get("why_it_matters"):
            result["why_it_matters"] = item["why_it_matters"]
        if item.get("potential_impact"):
            result["potential_impact"] = item["potential_impact"]
        if item.get("early_warning_signal"):
            result["early_warning"] = item["early_warning_signal"]
        if item.get("special_flags"):
            flags = item["special_flags"]
            result["special_flags"] = flags if isinstance(flags, list) else [str(flags)]
        if item.get("actors"):
            actors = item["actors"]
            result["actors"] = ", ".join(actors) if isinstance(actors, list) else str(actors)
        if item.get("attention_level") and item["attention_level"] != "Routine Monitoring":
            result["attention_level"] = item["attention_level"]
        return result
    
    key_developments = []
    added_ids = set()
    
    # Add ALL critical/high items (no cap - user wants comprehensive coverage)
    for item in deduped_critical:
        item_id = item.get("id")
        if item_id and item_id not in added_ids:
            if item.get("state") in NER_STATES + ["Multiple", "Bangladesh", "Myanmar", ""]:
                key_developments.append(build_brief_item(item))
                added_ids.add(item_id)
    
    # Add diverse NER items (medium+ severity)
    for item in diverse_ner_items:
        item_id = item.get("id")
        if item_id and item_id not in added_ids:
            key_developments.append(build_brief_item(item))
            added_ids.add(item_id)
    
    brief_data["key_developments"] = key_developments
    
    # ========== 9. BUILD NATIONAL NEWS ==========
    national_deduped = []
    for item in military_national:
        title = item.get("title", "")
        if item.get("id") not in added_ids and not is_duplicate_title(title, seen_titles, seen_entities):
            national_deduped.append(build_brief_item(item))
            seen_titles.append(normalize_title(title))
            seen_entities.append(extract_key_entities(title))
    brief_data["national_news"] = national_deduped[:15]
    
    # ========== 10. BUILD INTERNATIONAL NEWS ==========
    brief_data["international_news"] = [
        {
            **build_brief_item(item),
            "countries": ", ".join(item.get("countries_involved", [])) if isinstance(item.get("countries_involved"), list) else str(item.get("countries_involved", "")),
        }
        for item in diverse_intl_items
        if item.get("id") not in added_ids
    ]
    
    # ========== 11. ADD PATTERN INSIGHTS AND UPLOADS ==========
    # Build pattern insights for the brief
    brief_data["pattern_insights"] = [
        {
            "region": p.get("region", ""),
            "detail": p.get("detail", p.get("pattern_type", "")),
            "event_count": p.get("event_count", 0),
            "escalation_risk": p.get("escalation_risk", "LOW"),
            "avg_priority_score": p.get("avg_priority_score", 0),
            "window_days": p.get("window_days", 7),
            "sample_titles": p.get("sample_titles", [])[:2],
        }
        for p in detected_patterns
        if p.get("escalation_risk") in ("CRITICAL", "HIGH", "MODERATE")
    ][:15]
    
    brief_data["uploaded_insights"] = [
        {
            "filename": doc.get("filename", ""),
            "ai_analysis": doc.get("ai_analysis", doc.get("content_summary", "")),
            "region": doc.get("region", ""),
            "uploaded_at": doc.get("uploaded_at", "")
        }
        for doc in uploaded_docs[:10]
    ]
    
    # ========== 12. TRACK INCLUDED ITEM IDS ==========
    brief_data["included_item_ids"] = list(added_ids)
    
    # ========== 13. SAVE AND RETURN ==========
    brief = DailyBrief(**brief_data)
    doc = brief.model_dump()
    # Clear legacy fields not in the current model
    doc.pop("twitter_highlights", None)
    # Ensure clean upsert replacing old brief entirely
    await briefs_col.replace_one({"date": date}, doc, upsert=True)
    
    logger.info(f"Brief generated: {len(key_developments)} NER developments, {len(brief_data.get('national_news', []))} national, {len(brief_data.get('international_news', []))} international, {len(brief_data.get('pattern_insights', []))} patterns, {len(added_ids)} items tracked")
    
    return doc


def generate_manual_brief(items, date):
    developments = []
    state_highlights = {}
    cross_border_items = []

    for item in items[:15]:
        sev = item.get('severity', 'medium').upper()
        developments.append(f"[{sev}] {item['title']}")
        state = item.get('state', '')
        if state and state not in state_highlights:
            state_highlights[state] = item.get('ai_summary', item['title'])
        if item.get('is_cross_border'):
            cross_border_items.append(item['title'])

    critical_count = sum(1 for i in items if i.get('severity') == 'critical')
    high_count = sum(1 for i in items if i.get('severity') == 'high')

    return {
        "id": str(uuid.uuid4()),
        "date": date,
        "key_developments": developments[:8],
        "state_highlights": state_highlights,
        "cross_border_insights": "; ".join(cross_border_items[:3]) if cross_border_items else "No significant cross-border developments reported in this period.",
        "analyst_summary": f"Intelligence summary for {date}: {len(items)} items monitored across NER. {critical_count} critical and {high_count} high-severity items require immediate attention. Continuous monitoring of cross-border activities and insurgent movements recommended.",
        "generated_at": datetime.now(timezone.utc).isoformat()
    }


async def fetch_and_process_news(source_filter: str = None):
    """
    Fetch and process news with rate-limit-aware batching.
    
    source_filter: "grassroots", "standard", "established", or None (all)
    """
    from rss_fetcher import fetch_all_feeds, RSS_SOURCES, get_sources_by_priority
    
    # Configuration for rate limit management
    MAX_ARTICLES_PER_CYCLE = 25
    BATCH_SIZE = 3
    BATCH_PAUSE = 5
    INTER_ARTICLE_DELAY = 1.5
    
    # Filter sources if specified
    if source_filter:
        active_sources = get_sources_by_priority(source_filter)
        logger.info(f"Filtered to {len(active_sources)} {source_filter} sources")
    else:
        active_sources = RSS_SOURCES
    
    # Update scan status
    scan_status["is_scanning"] = True
    scan_status["progress"] = 0
    scan_status["total_sources"] = len(active_sources)
    scan_status["sources_scanned"] = 0
    scan_status["current_source"] = ""
    scan_status["articles_found"] = 0
    scan_status["relevant_found"] = 0
    scan_status["scan_log"] = []
    
    async def on_source_progress(idx, total, source_name):
        scan_status["sources_scanned"] = idx
        scan_status["current_source"] = source_name
        scan_status["progress"] = int((idx / total) * 100) if total > 0 else 0
        if source_name != "Complete":
            scan_status["scan_log"].append(source_name)
            # Keep only last 10 entries
            if len(scan_status["scan_log"]) > 10:
                scan_status["scan_log"] = scan_status["scan_log"][-10:]
    
    logger.info(f"=== Starting news fetch cycle ({source_filter or 'all'}) ===")
    try:
        # Step 1: Fetch RSS articles with progress tracking
        articles = await fetch_all_feeds(progress_callback=on_source_progress, sources=active_sources if source_filter else None)
        scan_status["articles_found"] = len(articles)
        logger.info(f"Fetched {len(articles)} relevant articles from RSS feeds")

        # Step 2: Deduplicate by URL AND title similarity
        import re
        def normalize_for_dedup(title):
            t = (title or "").lower().strip()
            t = re.sub(r'[^a-z0-9\s]', '', t)
            t = re.sub(r'\s+', ' ', t).strip()
            return t
        
        def title_is_similar(t1, existing_titles, threshold=0.65):
            words1 = set(t1.split())
            for t2 in existing_titles:
                words2 = set(t2.split())
                if not words2:
                    continue
                overlap = len(words1 & words2)
                total = max(len(words1), len(words2))
                if total > 0 and overlap / total >= threshold:
                    return True
            return False
        
        # Get recent titles from DB for title-level dedup
        recent_db_items = await intelligence_col.find(
            {"processed": True},
            {"title": 1, "source_url": 1, "_id": 0}
        ).sort("fetched_at", -1).limit(500).to_list(500)
        
        existing_urls = set(item.get("source_url", "") for item in recent_db_items)
        existing_titles = [normalize_for_dedup(item.get("title", "")) for item in recent_db_items]
        
        new_articles = []
        url_dupes = 0
        title_dupes = 0
        for article in articles:
            url = article.get("source_url", "")
            title = article.get("title", "")
            if not url:
                continue
            if url in existing_urls:
                url_dupes += 1
                continue
            norm_title = normalize_for_dedup(title)
            if len(norm_title) > 10 and title_is_similar(norm_title, existing_titles):
                title_dupes += 1
                continue
            new_articles.append(article)
            existing_urls.add(url)
            existing_titles.append(norm_title)
        
        skipped = url_dupes + title_dupes
        logger.info(f"Deduplication: {url_dupes} URL dupes, {title_dupes} title dupes, {len(new_articles)} new articles")

        # Step 2.5: Apply intelligence hard filter BEFORE AI processing
        from intelligence_filter import hard_filter, detect_language, run_filter_pipeline
        EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY', '')
        
        filtered_articles = []
        filter_rejected = 0
        translated_count = 0
        for article in new_articles:
            passed, reason = hard_filter(article)
            if not passed:
                filter_rejected += 1
                logger.debug(f"  Filtered out: {article.get('title', '')[:50]} ({reason})")
                continue
            
            # Detect language and translate if needed (pre-process for AI)
            lang = detect_language(f"{article.get('title', '')} {article.get('raw_content', '')[:200]}")
            article["detected_language"] = lang
            
            if lang != "en" and EMERGENT_LLM_KEY:
                try:
                    filter_result = await run_filter_pipeline(article, EMERGENT_LLM_KEY)
                    if filter_result.get("translated_title"):
                        article["original_title"] = article["title"]
                        article["title"] = filter_result["translated_title"]
                        translated_count += 1
                    if filter_result.get("translated_content"):
                        article["original_content"] = article.get("raw_content", "")
                        article["raw_content"] = filter_result["translated_content"]
                except Exception as e:
                    logger.warning(f"Translation failed, using original: {e}")
            
            filtered_articles.append(article)
        
        scan_status["filtered_out"] = filter_rejected
        scan_status["translated"] = translated_count
        scan_status["relevant_found"] = len(filtered_articles)
        logger.info(f"Intelligence Filter: {filter_rejected} rejected, {translated_count} translated, {len(filtered_articles)} passed")
        new_articles = filtered_articles

        if not new_articles:
            logger.info("No new articles to process after filtering. Cycle complete.")
            scan_status["is_scanning"] = False
            scan_status["progress"] = 100
            scan_status["current_source"] = ""
            scan_status["last_scan_at"] = datetime.now(timezone.utc).isoformat()
            scan_status["last_scan_result"] = {
                "feeds_scanned": len(RSS_SOURCES),
                "total_articles": len(articles),
                "new_relevant": 0,
                "duplicates_skipped": skipped,
                "filtered_out": filter_rejected,
                "translated": translated_count
            }
            return

        # Step 3: Limit to MAX_ARTICLES_PER_CYCLE
        if len(new_articles) > MAX_ARTICLES_PER_CYCLE:
            logger.info(f"Limiting to {MAX_ARTICLES_PER_CYCLE} articles this cycle")
            new_articles = new_articles[:MAX_ARTICLES_PER_CYCLE]

        # Step 4: Process new articles with Sifter tiered analysis
        from sifter import sift_article
        success_count = 0
        fail_count = 0
        skip_count = 0
        rate_limit_hits = 0

        for batch_start in range(0, len(new_articles), BATCH_SIZE):
            batch = new_articles[batch_start:batch_start + BATCH_SIZE]
            batch_num = (batch_start // BATCH_SIZE) + 1
            total_batches = (len(new_articles) + BATCH_SIZE - 1) // BATCH_SIZE
            
            logger.info(f"  Processing batch {batch_num}/{total_batches} ({len(batch)} articles)...")
            
            for article in batch:
                await asyncio.sleep(INTER_ARTICLE_DELAY)
                
                # Level 1 Sifter — fast pre-filter
                sift_result = sift_article(article)
                article["_sifter"] = sift_result
                
                result, was_rate_limited = await _classify_with_retry_v2(article)
                
                if was_rate_limited:
                    rate_limit_hits += 1
                
                if result is None:
                    raw_doc = _make_raw_doc(article)
                    await intelligence_col.insert_one(raw_doc)
                    fail_count += 1
                elif result.get("is_relevant", True):
                    # Apply Sifter boost to priority score
                    sifter_boost = sift_result.get("boost_score", 0)
                    result["priority_score"] = min(100, result.get("priority_score", 30) + sifter_boost)
                    result["sifter_level"] = sift_result["level"]
                    result["sifter_triggers"] = sift_result["triggers"]
                    
                    # Extract relationships for knowledge graph
                    if result.get("entities"):
                        ents = result["entities"]
                        relationships = []
                        for person in (ents.get("persons") or [])[:3]:
                            for loc in (ents.get("locations") or [])[:3]:
                                relationships.append({
                                    "actor": person, "location": loc,
                                    "context": result.get("threat_category", ""),
                                    "date": result.get("published_at", "")
                                })
                        for org in (ents.get("organizations") or [])[:3]:
                            for loc in (ents.get("locations") or [])[:3]:
                                relationships.append({
                                    "actor": org, "location": loc,
                                    "context": result.get("threat_category", ""),
                                    "date": result.get("published_at", "")
                                })
                        if relationships:
                            result["relationships"] = relationships[:10]
                    
                    item = IntelligenceItem(**result)
                    doc = item.model_dump()
                    await intelligence_col.insert_one(doc)
                    success_count += 1
                    invalidate_stats_cache()
                    
                    # Generate embedding async (don't block)
                    try:
                        from embedding_service import generate_embedding
                        emb_text = f"{doc.get('title', '')}. {doc.get('ai_summary', '')}"
                        emb = await generate_embedding(emb_text)
                        if emb:
                            await intelligence_col.update_one(
                                {"id": doc["id"]},
                                {"$set": {"embedding": emb}}
                            )
                    except Exception as emb_err:
                        logger.debug(f"Embedding generation skipped: {emb_err}")
                    
                    # Broadcast to WebSocket clients
                    try:
                        ws_msg = {
                            "type": "new_item",
                            "item": {
                                "id": doc.get("id"),
                                "title": doc.get("title"),
                                "severity": doc.get("severity"),
                                "priority_score": doc.get("priority_score", 0),
                                "confidence_score": doc.get("confidence_score", 70),
                                "threat_trajectory": doc.get("threat_trajectory", "INDETERMINATE"),
                                "state": doc.get("state"),
                                "source": doc.get("source"),
                                "threat_category": doc.get("threat_category"),
                                "published_at": doc.get("published_at"),
                                "ai_summary": doc.get("ai_summary", "")[:200],
                            }
                        }
                        # Sifter L2 items or critical/high = elite alert
                        if sift_result["level"] == 2 or doc.get("severity") in ("critical", "high"):
                            ws_msg["type"] = "elite_alert"
                            ws_msg["sifter_triggers"] = sift_result["triggers"]
                        await ws_manager.broadcast(ws_msg)
                    except Exception as ws_err:
                        logger.debug(f"WS broadcast error: {ws_err}")
                else:
                    skip_count += 1

            if batch_start + BATCH_SIZE < len(new_articles):
                pause_time = BATCH_PAUSE * 2 if rate_limit_hits > 2 else BATCH_PAUSE
                logger.info(f"  Batch {batch_num} done. Success: {success_count}, Failed: {fail_count}. "
                            f"Pausing {pause_time}s...")
                await asyncio.sleep(pause_time)

        logger.info("=== Fetch cycle complete ===")
        logger.info(f"  Processed: {success_count} | Failed: {fail_count} | Not relevant: {skip_count}")
        logger.info(f"  Duplicates skipped: {skipped} | Filtered out: {filter_rejected} | Translated: {translated_count}")
        logger.info(f"  Rate limit hits: {rate_limit_hits}")

        # Finalize scan status
        scan_status["is_scanning"] = False
        scan_status["progress"] = 100
        scan_status["current_source"] = ""
        scan_status["last_scan_at"] = datetime.now(timezone.utc).isoformat()
        scan_status["last_scan_result"] = {
            "feeds_scanned": scan_status["total_sources"],
            "total_articles": len(articles),
            "new_relevant": success_count,
            "duplicates_skipped": skipped,
            "filtered_out": filter_rejected,
            "translated": translated_count,
            "failed": fail_count,
            "not_relevant": skip_count
        }

        # Trigger pattern detection after processing
        try:
            from pattern_engine import detect_patterns
            asyncio.create_task(detect_patterns(db))
        except Exception as e:
            logger.warning(f"Pattern detection trigger failed: {e}")
        
        # Broadcast scan complete to WebSocket clients
        try:
            await ws_manager.broadcast({
                "type": "scan_complete",
                "result": scan_status["last_scan_result"]
            })
        except Exception:
            pass

    except Exception as e:
        logger.error(f"News fetch cycle failed: {e}")
        scan_status["is_scanning"] = False
        scan_status["progress"] = 100
        scan_status["current_source"] = ""
        scan_status["last_scan_at"] = datetime.now(timezone.utc).isoformat()
        scan_status["last_scan_result"] = {"error": str(e)}


async def _classify_with_retry(article, max_retries=3):
    """Classify an article with exponential backoff retry on failure."""
    result, _ = await _classify_with_retry_v2(article, max_retries)
    return result


async def _classify_with_retry_v2(article, max_retries=4):
    """
    Classify an article with aggressive exponential backoff for rate limits.
    
    Returns: (result, was_rate_limited)
    - result: classification dict or None if all retries failed
    - was_rate_limited: True if we hit rate limits during processing
    
    Backoff strategy:
    - Normal errors: 3s, 6s, 12s, 24s (doubling)
    - Rate limit errors: 15s, 30s, 60s, 120s (aggressive)
    """
    RATE_LIMIT_INDICATORS = ['rate', '429', 'limit', 'quota', 'too many', 'throttle']
    was_rate_limited = False
    
    for attempt in range(max_retries):
        try:
            result = await asyncio.wait_for(
                asyncio.to_thread(_sync_classify, article),
                timeout=60  # Increased timeout to 60s
            )
            return result, was_rate_limited  # Success
            
        except asyncio.TimeoutError:
            base_wait = 5
            wait = base_wait * (2 ** attempt)  # 5s, 10s, 20s, 40s
            logger.warning(f"  [Attempt {attempt + 1}/{max_retries}] Timeout for: "
                           f"{article.get('title', '')[:40]}... retrying in {wait}s")
            await asyncio.sleep(wait)
            
        except Exception as e:
            err_str = str(e).lower()
            is_rate_limit = any(indicator in err_str for indicator in RATE_LIMIT_INDICATORS)
            
            if is_rate_limit:
                was_rate_limited = True
                # Aggressive backoff for rate limits: 15s, 30s, 60s, 120s
                wait = 15 * (2 ** attempt)
                logger.warning(f"  [Attempt {attempt + 1}/{max_retries}] RATE LIMIT for: "
                               f"{article.get('title', '')[:40]}... backing off {wait}s")
            else:
                # Normal backoff for other errors: 3s, 6s, 12s, 24s
                wait = 3 * (2 ** attempt)
                logger.warning(f"  [Attempt {attempt + 1}/{max_retries}] Error ({str(e)[:60]}) for: "
                               f"{article.get('title', '')[:40]}... retrying in {wait}s")
            
            await asyncio.sleep(wait)
    
    logger.error(f"  All {max_retries} retries exhausted for: {article.get('title', '')[:50]}")
    return None, was_rate_limited  # Signal failure


def _sync_classify(article):
    """Synchronous wrapper for AI classification to run in thread pool"""
    import asyncio
    loop = asyncio.new_event_loop()
    try:
        from ai_pipeline import classify_and_analyze_article
        return loop.run_until_complete(classify_and_analyze_article(article))
    finally:
        loop.close()


def _make_raw_doc(article):
    return {
        "id": str(uuid.uuid4()),
        "title": article.get("title", ""),
        "source": article.get("source", "Unknown"),
        "source_url": article.get("source_url", ""),
        "published_at": article.get("published_at", datetime.now(timezone.utc).isoformat()),
        "fetched_at": datetime.now(timezone.utc).isoformat(),
        "raw_content": article.get("raw_content", "")[:5000],
        "ai_summary": article.get("raw_content", "")[:200],
        "why_it_matters": "Pending AI analysis.",
        "potential_impact": "Assessment pending.",
        "attention_level": "Monitor",
        "state": "",
        "threat_category": "",
        "severity": "low",
        "is_cross_border": article.get("region", "") in ("Bangladesh", "Myanmar"),
        "countries_involved": [article["region"]] if article.get("region") in ("Bangladesh", "Myanmar") else [],
        "processed": False,
        "tags": ["unprocessed"]
    }


async def analyze_uploaded_document(doc_id: str):
    """Analyze an uploaded document using AI"""
    doc = await uploads_col.find_one({"id": doc_id}, {"_id": 0})
    if not doc:
        logger.error(f"Document {doc_id} not found")
        return
    
    extracted_text = doc.get("extracted_text", "")
    if not extracted_text:
        logger.warning(f"No text extracted from document {doc_id}")
        return
    
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        
        EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY', '')
        
        analysis_prompt = """You are a military intelligence analyst. Analyze this document and provide:
1. A concise summary (3-4 lines)
2. Key intelligence points relevant to India's North Eastern Region, Bangladesh, or Myanmar
3. Any security implications
4. Recommended attention level (Immediate Action Required, Priority Monitoring, Active Monitoring, Monitor)
5. Primary region affected (if identifiable): Assam, Meghalaya, Mizoram, Manipur, Arunachal Pradesh, Tripura, Bangladesh, Myanmar, or National/International

Respond in JSON format:
{
  "summary": "...",
  "key_points": ["...", "..."],
  "security_implications": "...",
  "attention_level": "...",
  "region": "..."
}"""
        
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"doc-{doc_id}",
            system_message=analysis_prompt
        ).with_model("anthropic", "claude-haiku-4-5-20251001")
        
        user_message = UserMessage(text=f"Analyze this document:\n\n{extracted_text[:4000]}")
        response = await chat.send_message(user_message)
        
        response_text = str(response)
        
        # Parse JSON from response
        import json
        json_start = response_text.find('{')
        json_end = response_text.rfind('}') + 1
        if json_start >= 0 and json_end > json_start:
            analysis = json.loads(response_text[json_start:json_end])
        else:
            analysis = {"summary": response_text[:500], "region": ""}
        
        # Update document with AI analysis
        await uploads_col.update_one(
            {"id": doc_id},
            {"$set": {
                "ai_analysis": analysis.get("summary", "") + "\n\nKey Points:\n" + "\n".join(analysis.get("key_points", [])),
                "region": analysis.get("region", ""),
                "processed": True
            }}
        )
        
        logger.info(f"Successfully analyzed document {doc_id}")
        
    except Exception as e:
        logger.error(f"Error analyzing document {doc_id}: {e}")
        await uploads_col.update_one(
            {"id": doc_id},
            {"$set": {"ai_analysis": f"Analysis failed: {str(e)}", "processed": True}}
        )


async def bulk_scrape_all_feeds():
    """
    BULK SCRAPE: Fetch ALL articles from ALL RSS feeds and store them as unprocessed.
    
    This is the initial data collection phase - no AI processing.
    Articles will be gradually AI-processed by the retry scheduler.
    
    Benefits:
    - Collects maximum articles quickly (no rate limit concerns)
    - AI processing happens gradually via scheduled retries
    - Keeps LLM costs low by spreading processing over time
    """
    from rss_fetcher import fetch_all_feeds
    
    logger.info("=" * 60)
    logger.info("=== BULK SCRAPE: Fetching ALL articles from RSS feeds ===")
    logger.info("=" * 60)
    
    try:
        # Step 1: Fetch all RSS articles
        articles = await fetch_all_feeds()
        logger.info(f"Fetched {len(articles)} total articles from RSS feeds")

        # Step 2: Deduplicate — only keep articles NOT already in DB
        new_articles = []
        for article in articles:
            url = article.get("source_url", "")
            if not url:
                continue
            existing = await intelligence_col.find_one({"source_url": url}, {"_id": 1})
            if not existing:
                new_articles.append(article)
        
        skipped = len(articles) - len(new_articles)
        logger.info(f"Deduplication: {skipped} already in DB, {len(new_articles)} new articles to store")

        if not new_articles:
            logger.info("No new articles to store. Bulk scrape complete.")
            return {"stored": 0, "skipped": skipped, "total_fetched": len(articles)}

        # Step 3: Store ALL new articles as unprocessed (NO AI processing)
        stored_count = 0
        for article in new_articles:
            raw_doc = _make_raw_doc(article)
            await intelligence_col.insert_one(raw_doc)
            stored_count += 1
            
            # Log progress every 50 articles
            if stored_count % 50 == 0:
                logger.info(f"  Stored {stored_count}/{len(new_articles)} articles...")

        logger.info("=" * 60)
        logger.info("=== BULK SCRAPE COMPLETE ===")
        logger.info(f"  Stored: {stored_count} new articles (unprocessed)")
        logger.info(f"  Skipped: {skipped} duplicates")
        logger.info(f"  Total in DB: {await intelligence_col.count_documents({})}")
        logger.info(f"  Pending AI processing: {await intelligence_col.count_documents({'processed': False})}")
        logger.info("=" * 60)
        logger.info("Articles will be AI-processed gradually via scheduled retries (every 15 min)")
        
        return {"stored": stored_count, "skipped": skipped, "total_fetched": len(articles)}

    except Exception as e:
        logger.error(f"Bulk scrape failed: {e}")
        raise


async def analyze_unprocessed_items():
    """
    Retry AI classification on previously failed items with exponential backoff.
    
    This runs every 15 minutes and processes items that failed in previous cycles.
    Uses conservative limits to avoid overwhelming the API.
    """
    # Configuration for retry processing - increased for faster processing
    MAX_RETRY_PER_CYCLE = 20  # Process max 20 unprocessed items per retry cycle
    INTER_ARTICLE_DELAY = 2.5  # 2.5 seconds between articles for safety
    
    unprocessed = await intelligence_col.find(
        {"processed": False}, {"_id": 0}
    ).limit(MAX_RETRY_PER_CYCLE).to_list(MAX_RETRY_PER_CYCLE)

    if not unprocessed:
        logger.info("No unprocessed items to retry.")
        return

    total_unprocessed = await intelligence_col.count_documents({"processed": False})
    logger.info(f"=== AI Processing Cycle: {len(unprocessed)}/{total_unprocessed} unprocessed items ===")
    
    success = 0
    not_relevant = 0
    rate_limit_hits = 0
    
    for idx, item in enumerate(unprocessed):
        await asyncio.sleep(INTER_ARTICLE_DELAY)
        result, was_rate_limited = await _classify_with_retry_v2(item, max_retries=3)
        
        if was_rate_limited:
            rate_limit_hits += 1
            # If we hit too many rate limits, stop early and wait for next cycle
            if rate_limit_hits >= 3:
                logger.warning(f"  Hit {rate_limit_hits} rate limits, stopping retry cycle early. "
                               f"Will continue in next cycle.")
                break
        
        if result and result.get("is_relevant", True):
            update_fields = {k: v for k, v in result.items() if k != "_id"}
            update_fields["processed"] = True
            update_fields.pop("is_relevant", None)
            await intelligence_col.update_one(
                {"id": item["id"]},
                {"$set": update_fields}
            )
            success += 1
        elif result and not result.get("is_relevant", True):
            # Mark as processed but not relevant
            await intelligence_col.update_one(
                {"id": item["id"]},
                {"$set": {"processed": True, "tags": ["not_relevant"]}}
            )
            not_relevant += 1
        
        # Log progress every 5 articles
        if (idx + 1) % 5 == 0:
            logger.info(f"  Progress: {idx + 1}/{len(unprocessed)} | Success: {success} | Not relevant: {not_relevant}")
    
    remaining = total_unprocessed - success - not_relevant
    logger.info("=== AI Processing Complete ===")
    logger.info(f"  Processed: {success} relevant | {not_relevant} not relevant")
    logger.info(f"  Remaining unprocessed: {remaining}")
    logger.info(f"  Rate limit hits: {rate_limit_hits}")


async def initialize_sources():
    """Seed RSS sources from rss_fetcher config — always update to latest list"""
    from rss_fetcher import RSS_SOURCES
    # Upsert all sources to ensure new ones are added
    for source in RSS_SOURCES:
        await sources_col.update_one(
            {"url": source["url"]},
            {"$set": {**source, "id": str(uuid.uuid4())}},
            upsert=True
        )
    count = await sources_col.count_documents({})
    logger.info(f"RSS sources synced: {count} total ({len(RSS_SOURCES)} configured)")


async def generate_scheduled_daily_brief():
    """Auto-generate the daily brief at 0600 IST.
    Called by APScheduler cron trigger."""
    import pytz
    ist = pytz.timezone("Asia/Kolkata")
    now_ist = datetime.now(timezone.utc).astimezone(ist)
    date = now_ist.strftime("%Y-%m-%d")
    logger.info(f"=== SCHEDULED DAILY BRIEF GENERATION: {date} at {now_ist.strftime('%H:%M IST')} ===")
    try:
        brief = await generate_brief_for_date(date)
        dev_count = len(brief.get("key_developments", []))
        pattern_count = len(brief.get("pattern_insights", []))
        logger.info(f"Scheduled brief generated: {dev_count} developments, {pattern_count} pattern insights")
    except Exception as e:
        logger.error(f"Scheduled brief generation failed: {e}")


async def fetch_grassroots_sources():
    """Fetch only grassroots NER/cross-border sources (high priority, 60min cycle)."""
    from rss_fetcher import get_sources_by_priority
    grassroots = get_sources_by_priority("grassroots")
    logger.info(f"=== Grassroots fetch: {len(grassroots)} sources ===")
    await fetch_and_process_news(source_filter="grassroots")


async def fetch_established_sources():
    """Fetch established/government security databases (12hr cycle)."""
    from rss_fetcher import get_sources_by_priority
    established = get_sources_by_priority("established")
    logger.info(f"=== Established fetch: {len(established)} sources ===")
    await fetch_and_process_news(source_filter="established")


async def run_embedding_backfill():
    """Periodic embedding backfill for items without embeddings."""
    try:
        from embedding_service import backfill_embeddings
        count = await backfill_embeddings(db, 30)
        logger.info(f"Embedding backfill: {count} items processed")
    except Exception as e:
        logger.warning(f"Embedding backfill failed: {e}")


@app.on_event("startup")
async def startup():
    await initialize_sources()
    # Only trigger initial fetch if DB is empty (fresh start)
    item_count = await intelligence_col.count_documents({})
    if item_count == 0:
        logger.info("Empty database — triggering initial fetch...")
        asyncio.create_task(fetch_and_process_news())
    else:
        logger.info(f"Database has {item_count} items. Skipping startup fetch (scheduler will handle next cycle).")
        # Retry any unprocessed items from previous cycles
        unprocessed = await intelligence_col.count_documents({"processed": False})
        if unprocessed > 0:
            logger.info(f"{unprocessed} unprocessed items found — triggering retry...")
            asyncio.create_task(analyze_unprocessed_items())

    try:
        from apscheduler.schedulers.asyncio import AsyncIOScheduler
        from apscheduler.triggers.cron import CronTrigger
        scheduler = AsyncIOScheduler()
        # Grassroots NER/cross-border sites: every 60 min
        scheduler.add_job(fetch_grassroots_sources, 'interval', minutes=60, id='grassroots_fetch')
        # Standard sources (national, intl): every 30 min
        scheduler.add_job(fetch_and_process_news, 'interval', minutes=30, id='news_fetch')
        # Established security databases: every 12 hours
        scheduler.add_job(fetch_established_sources, 'interval', hours=12, id='established_fetch')
        # Retry unprocessed
        scheduler.add_job(analyze_unprocessed_items, 'interval', minutes=15, id='retry_unprocessed')
        # Daily brief auto-generation at 0600 IST (0030 UTC)
        scheduler.add_job(
            generate_scheduled_daily_brief,
            CronTrigger(hour=0, minute=30, timezone='UTC'),
            id='daily_brief_0600',
            misfire_grace_time=3600
        )
        # Embedding backfill every 6 hours
        scheduler.add_job(run_embedding_backfill, 'interval', hours=6, id='embedding_backfill')
        scheduler.start()
        logger.info("Scheduler: grassroots/60min, standard/30min, established/12hr, retry/15min, brief/0600 IST, embeddings/6hr")
    except Exception as e:
        logger.warning(f"Scheduler setup failed: {e}")


@app.on_event("shutdown")
async def shutdown():
    client.close()


# ============================================================
# WebSocket Endpoint (mounted on app directly, not api_router)
# ============================================================
@app.websocket("/api/ws/intelligence")
async def websocket_intelligence(websocket: WebSocket):
    """Real-time intelligence feed via WebSocket.
    Sends new items and critical alerts as they're processed."""
    await ws_manager.connect(websocket)
    try:
        while True:
            # Keep connection alive, listen for client pings
            data = await websocket.receive_text()
            if data == "ping":
                await websocket.send_json({"type": "pong"})
    except WebSocketDisconnect:
        ws_manager.disconnect(websocket)
    except Exception:
        ws_manager.disconnect(websocket)


app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)
